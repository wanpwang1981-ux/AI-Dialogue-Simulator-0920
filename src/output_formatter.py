import csv
import io
from typing import List, Dict
from docx import Document
from openpyxl import Workbook

# 定義一個標準的對話紀錄結構
# 這是從主應用傳遞到格式化器的資料格式
StructuredLog = List[Dict[str, str]]

def to_txt(log: StructuredLog) -> str:
    """
    將結構化日誌轉換為純文字格式。

    Args:
        log (StructuredLog): 結構化的對話日誌。

    Returns:
        str: 格式化後的純文字字串。
    """
    output = []
    for entry in log:
        # 如果是系統訊息(例如開頭的角色介紹)，直接印出內容
        if entry['speaker'] == 'System':
            output.append(entry['content'])
        else:
            output.append(f"[{entry['speaker']}]:\n{entry['content']}\n")
    return "\n".join(output)

def to_csv(log: StructuredLog) -> str:
    """
    將結構化日誌轉換為CSV格式字串。

    Args:
        log (StructuredLog): 結構化的對話日誌。

    Returns:
        str: CSV格式的字串。
    """
    # 過濾掉非對話的系統訊息
    dialogue_only_log = [entry for entry in log if entry['speaker'] != 'System']

    # 使用 io.StringIO 在記憶體中建立一個類似檔案的物件來寫入CSV
    output = io.StringIO()
    # 定義CSV的欄位標頭
    fieldnames = ['speaker', 'content']
    writer = csv.DictWriter(output, fieldnames=fieldnames)

    writer.writeheader()
    writer.writerows(dialogue_only_log)

    # 獲取StringIO物件中的完整字串內容
    return output.getvalue()

def to_md(log: StructuredLog) -> str:
    """
    將結構化日誌轉換為Markdown格式。

    Args:
        log (StructuredLog): 結構化的對話日誌。

    Returns:
        str: Markdown格式的字串。
    """
    output = []
    for entry in log:
        if entry['speaker'] == 'System':
            # 系統訊息直接當作前言
            output.append(entry['content'])
        else:
            # 使用 '###' 作為每個發言者的標題
            # 使用 '>' 來引用發言內容
            speaker = entry['speaker']
            content = entry['content'].replace('\n', '\n> ') # 處理多行內容
            output.append(f"### {speaker}")
            output.append(f"> {content}\n")
    return "\n".join(output)

def to_docx(log: StructuredLog, filepath: str):
    """
    將結構化日誌寫入一個Word (.docx) 檔案。

    Args:
        log (StructuredLog): 結構化的對話日誌。
        filepath (str): 要儲存的檔案路徑。
    """
    document = Document()
    # 處理開頭的系統訊息
    if log and log[0]['speaker'] == 'System':
        # 簡單地將介紹文字加入
        intro_text = log[0]['content']
        document.add_paragraph(intro_text)
        # 過濾掉介紹，只留下對話
        dialogue_only_log = log[1:]
    else:
        dialogue_only_log = log

    document.add_heading('對話紀錄', level=1)
    for entry in dialogue_only_log:
        document.add_heading(entry['speaker'], level=3)
        document.add_paragraph(entry['content'])
        document.add_paragraph() # 增加一些間距
    document.save(filepath)

def to_xlsx(log: StructuredLog, filepath: str):
    """
    將結構化日誌寫入一個Excel (.xlsx) 檔案。

    Args:
        log (StructuredLog): 結構化的對話日誌。
        filepath (str): 要儲存的檔案路徑。
    """
    # 過濾掉非對話的系統訊息
    dialogue_only_log = [entry for entry in log if entry['speaker'] != 'System']

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "對話紀錄"

    # 寫入標頭
    sheet['A1'] = "speaker"
    sheet['B1'] = "content"

    # 寫入資料
    for row_idx, entry in enumerate(dialogue_only_log, start=2):
        sheet[f'A{row_idx}'] = entry['speaker']
        sheet[f'B{row_idx}'] = entry['content']

    workbook.save(filepath)
